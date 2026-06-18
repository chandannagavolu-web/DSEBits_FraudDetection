# -*- coding: utf-8 -*-
"""
Generate the Mid-Semester Progress Report (.docx) for the M.Tech dissertation:
"Graph Neural Networks for Systemic Risk and Fraud Detection in Credit Systems".

All prose is written fresh (original) to avoid plagiarism. Run with:
    python build_midsem_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3B, 0x73)
ACCENT = RGBColor(0x2E, 0x5E, 0x8C)

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for hname, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
    st = doc.styles[hname]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = NAVY
    st.font.bold = True


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_para(text="", size=11, bold=False, italic=False, align=None,
             color=None, space_after=6, space_before=0, line=1.15):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def body(text, justify=True):
    p = add_para(text, size=11, line=1.4, space_after=8)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    p.add_run(text)
    return p


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


# =====================================================================
# TITLE PAGE
# =====================================================================
add_para("Mid-Semester Progress Report", size=20, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY, space_before=40, space_after=4)
add_para("M.Tech Dissertation (DSECLZG628T)", size=12, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para("GRAPH NEURAL NETWORKS FOR SYSTEMIC RISK AND FRAUD DETECTION IN CREDIT SYSTEMS",
         size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY, space_after=40)

meta = [
    ("Course No.", "DSECLZG628T"),
    ("Course Title", "Dissertation"),
    ("Student Name", "Chandan Nagavolu"),
    ("BITS ID", "2024da04063"),
    ("Degree Program", "M.Tech"),
    ("Research Area", "Credit Risk Management"),
    ("Dissertation carried out at", "Infosys Ltd, Hyderabad"),
    ("Supervisor", "Ravi Kiran Kesanakurti (Senior Data Scientist, Infosys Ltd)"),
    ("Additional Examiner", "Pratap Gottumukkala (Senior Project Manager, Infosys Ltd)"),
    ("Report Period", "May 2026 - June 2026 (Mid-Semester)"),
]
t = doc.add_table(rows=0, cols=2)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for k, v in meta:
    row = t.add_row().cells
    rp = row[0].paragraphs[0]
    rr = rp.add_run(k)
    rr.bold = True
    rr.font.size = Pt(11)
    vp = row[1].paragraphs[0]
    vr = vp.add_run(v)
    vr.font.size = Pt(11)
for r in t.rows:
    r.cells[0].width = Inches(2.4)
    r.cells[1].width = Inches(4.0)

add_para("", space_after=24)
add_para("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY, space_after=2)
add_para("Work Integrated Learning Programmes (WILP) Division", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Vidya Vihar, Pilani, Rajasthan - 333031", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Second Semester, Academic Year 2025-2026", size=11, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# =====================================================================
# ABSTRACT
# =====================================================================
h1("Abstract")
body(
    "Traditional credit-risk and fraud-detection models score each customer or transaction in "
    "isolation, ignoring the relational and temporal structure that fraud rings, money-laundering "
    "chains and contagion effects rely on. This dissertation investigates whether deep-learning "
    "models that explicitly encode these structures can outperform established machine-learning "
    "baselines. It applies three complementary families - Graph Neural Networks (GNNs) for "
    "interconnected fraud and systemic risk, Temporal Convolutional Networks (TCNs) for "
    "multi-period default prediction, and Autoencoders for unsupervised early-warning anomaly "
    "detection - and benchmarks each against traditional baselines (Logistic Regression, Random "
    "Forest, gradient-boosted trees) using standard metrics (Accuracy, Precision, Recall, F1, "
    "ROC-AUC) and industry risk metrics (PR-AUC, MCC, KS, Gini). This mid-semester report covers "
    "progress through the Literature Review and Design & Development phases - dataset selection, "
    "system architecture, graph construction, the evaluation harness, baseline and initial GNN "
    "implementations, and preliminary findings - and concludes with the plan for the remaining "
    "Testing, Review and Submission phases."
)

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS (Word field, updates on open)
# =====================================================================
h1("Table of Contents")
p = doc.add_paragraph()
run = p.add_run()
fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = r'TOC \o "1-3" \h \z \u'
fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
hint = OxmlElement("w:t"); hint.text = "Right-click and choose 'Update Field' to build the table of contents."
fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep)
run._r.append(hint); run._r.append(fldEnd)

doc.add_page_break()

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
h1("1. Introduction")

h2("1.1 Broad Area of Work")
body(
    "The broad area of this dissertation is the application of deep learning to credit-risk "
    "management, with a specific focus on Graph Neural Networks (GNNs) for risk and fraud "
    "detection. The study designs models that operate on graph-based representations of financial "
    "data, builds equivalent traditional machine-learning models for comparison, and evaluates "
    "the two families head-to-head to determine whether the graph-based approach delivers a "
    "measurable advantage in detecting fraud and predicting default."
)

h2("1.2 Background and Motivation")
body(
    "The expansion of digital payments, cross-border settlement and platform-based lending has "
    "made the modern financial landscape extremely complex. Entities such as accounts, cards, "
    "devices and merchants interact through millions of transactions, and adversarial behaviour "
    "increasingly hides inside these interactions rather than inside any single record. "
    "Organised fraud is rarely the act of one isolated account; it manifests as coordinated rings, "
    "layering chains and shared infrastructure. Likewise, systemic credit risk propagates through "
    "the relationships between counterparties, so the failure of one obligor can stress others."
)
body(
    "Conventional statistical and machine-learning scorecards are built almost entirely on "
    "structured, per-entity features. They are effective and interpretable, but by design they "
    "ignore two signals that are central to financial risk: the relational structure that links "
    "entities, and the temporal pattern of how behaviour changes over time. Recent advances in "
    "deep learning offer tools that target exactly these gaps. GNNs propagate information along "
    "the edges of a graph so that a node's prediction is informed by its neighbourhood; TCNs model "
    "ordered sequences to capture behavioural drift across periods; and autoencoders learn a "
    "compact representation of 'normal' behaviour so that deviations can be flagged without "
    "labelled fraud examples."
)

h2("1.3 Problem Statement")
body(
    "Given publicly available or simulated financial datasets, the problem is to determine whether "
    "models that explicitly encode relational and temporal structure - GNNs, TCNs and autoencoders "
    "- can improve fraud detection and credit-default prediction relative to strong traditional "
    "baselines, when both are evaluated under a single, consistent evaluation harness using metrics "
    "that are meaningful to risk practitioners."
)

h2("1.4 Objectives")
for o in [
    "Design and implement GNN models for detecting interconnected fraud and systemic risk, such as "
    "fraud rings and contagion across counterparties.",
    "Develop temporal deep-learning models (TCNs) for multi-period credit-default prediction.",
    "Build autoencoder-based anomaly-detection systems to generate early-warning signals.",
    "Benchmark the deep-learning models against traditional ML baselines (Logistic Regression, "
    "Random Forest and gradient-boosted trees) on a common evaluation harness.",
    "Integrate structured (tabular), relational (graph) and temporal (sequence) data into a unified "
    "predictive framework.",
]:
    bullet(o)

h2("1.5 Scope and Constraints")
body(
    "The study is restricted to publicly available or synthetically simulated financial datasets; "
    "no proprietary or real customer data is used, in keeping with privacy constraints. "
    "Implementation uses Python with PyTorch and PyTorch Geometric for graph models. The work is "
    "organised into two parallel tracks: a fraud / systemic-risk track that is graph-native and "
    "GNN-centric, and a credit-risk track that is tabular and temporal, scored with credit-scoring "
    "metrics. The emphasis is on a rigorous, reproducible comparison rather than on production "
    "deployment."
)

# =====================================================================
# 2. LITERATURE REVIEW
# =====================================================================
h1("2. Literature Review")
body(
    "The literature review surveyed three threads relevant to the dissertation: graph-based fraud "
    "detection, temporal models for credit default, and unsupervised anomaly detection. The "
    "findings below summarise the state of the art and identify the gap this work addresses."
)

h2("2.1 Graph Neural Networks for Financial Fraud")
body(
    "A consistent theme across recent surveys is that fraud is fundamentally relational. Cheng, "
    "Zou, Xiang and Jiang, in their review of GNNs for financial fraud detection, organise the "
    "field around how transaction and entity graphs are constructed and how message-passing "
    "architectures - graph convolutional networks (GCN), graph attention networks (GAT) and "
    "GraphSAGE - are adapted to the severe class imbalance and the presence of camouflaged "
    "fraudsters typical of financial graphs. Polu and colleagues report on modelling financial "
    "transaction networks at scale, emphasising the engineering needed to construct heterogeneous "
    "graphs from raw transaction logs and to train GNNs on graphs with very large node counts. "
    "The common conclusion is that representing data as a graph and allowing a node's "
    "classification to depend on its neighbourhood captures coordinated fraud patterns that "
    "per-record models miss."
)

h2("2.2 Temporal Models for Credit Default")
body(
    "Credit default is an inherently temporal phenomenon: repayment behaviour, utilisation and "
    "delinquency evolve over many periods before default crystallises. Recurrent architectures "
    "have traditionally been used for such sequences, but Temporal Convolutional Networks have "
    "emerged as a strong alternative, using dilated causal convolutions to capture long-range "
    "dependencies with stable, parallelisable training. For multi-period default prediction, this "
    "allows the model to learn how a borrower's trajectory - rather than a single snapshot - "
    "relates to eventual default."
)

h2("2.3 Autoencoders for Anomaly Detection")
body(
    "Because confirmed fraud labels are scarce and arrive late, unsupervised methods are valuable "
    "for early warning. Autoencoders learn to reconstruct normal behaviour; transactions or "
    "accounts with high reconstruction error are treated as anomalies. This approach is "
    "well-suited to the extreme class imbalance of fraud data and complements supervised models "
    "by flagging novel patterns that were not present in the training labels."
)

h2("2.4 Research Gap")
body(
    "While each modelling family is individually well studied, comparative studies that place "
    "GNNs, temporal models and autoencoders against strong traditional baselines under a single, "
    "consistent evaluation harness - and that report the risk-industry metrics practitioners "
    "actually use (PR-AUC, MCC, KS and Gini) rather than accuracy alone - remain limited. This "
    "dissertation addresses that gap by integrating structured, relational and temporal data into "
    "one framework and comparing the approaches fairly."
)

# =====================================================================
# 3. METHODOLOGY
# =====================================================================
h1("3. Proposed Methodology")

h2("3.1 Datasets")
body(
    "Datasets were selected to cover both project tracks and to satisfy the public/simulated-data "
    "constraint. The table below lists the candidates evaluated and their role."
)

ds = [
    ("Dataset", "Track", "Role in the study"),
    ("Elliptic Bitcoin", "Fraud (graph-native)", "~200K nodes / 234K edges of Bitcoin transactions labelled licit/illicit/unknown; primary GNN node-classification benchmark."),
    ("IEEE-CIS Fraud", "Fraud (graph construction)", "~590K e-commerce transactions with identity/device fields linked into a heterogeneous graph (card, device, email, address)."),
    ("PaySim (synthetic)", "Fraud / systemic", "Simulated mobile-money transfers forming a sender-receiver graph; useful for fraud-ring and laundering patterns."),
    ("German Credit", "Credit (tabular)", "1,000 applicants with good/bad labels; small, well-understood baseline for KS and Gini."),
    ("Credit Card Fraud (ULB)", "Fraud (tabular)", "~284K highly imbalanced transactions; tabular baseline and autoencoder anomaly benchmark."),
]
tbl = doc.add_table(rows=0, cols=3)
tbl.style = "Light Grid Accent 1"
for i, (a, b, c) in enumerate(ds):
    cells = tbl.add_row().cells
    for cell, txt in zip(cells, (a, b, c)):
        para = cell.paragraphs[0]
        run = para.add_run(txt)
        run.font.size = Pt(9.5)
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, "1F3B73")
add_para("", space_after=4)
body(
    "Final selection focuses the GNN/fraud track on the Elliptic dataset (graph-native, well "
    "labelled) with IEEE-CIS used for heterogeneous-graph construction, and the credit track on the "
    "German Credit dataset for scorecard metrics, with the ULB Credit-Card-Fraud set used for the "
    "autoencoder benchmark. Sources, licences and preprocessing notes are recorded in data/README.md."
)

h2("3.2 Graph Construction")
body(
    "For graph-native data (Elliptic), nodes are transactions and edges are the payment flows "
    "between them; each node carries its provided feature vector and a temporal step. For tabular "
    "transaction data (IEEE-CIS, PaySim), a heterogeneous graph is constructed by linking "
    "transactions to the entities they share - card, device, email domain and address - so that "
    "transactions connected through common infrastructure become neighbours. This converts "
    "implicit relationships in the raw tables into explicit edges that message passing can exploit."
)

h2("3.3 Model Families")
bullet("Baselines: Logistic Regression, Random Forest, and a gradient-boosted tree model "
       "(XGBoost/LightGBM) trained on engineered tabular features.")
bullet("GNNs: GCN, GraphSAGE and GAT for node classification on the constructed graphs, with "
       "class-imbalance handling (weighted loss / sampling).")
bullet("Temporal: a TCN with dilated causal convolutions for multi-period default prediction.")
bullet("Autoencoder: an undercomplete autoencoder trained on normal transactions, using "
       "reconstruction error as the anomaly score.")

h2("3.4 Evaluation Framework")
body(
    "All models are scored through a single shared evaluation module (src/evaluation/metrics.py) so "
    "that results are directly comparable. Alongside the standard metrics (Accuracy, Precision, "
    "Recall, F1 and ROC-AUC), the harness reports the risk-industry metrics defined for each track:"
)
bullet("Fraud track: PR-AUC (area under the precision-recall curve, preferred under heavy class "
       "imbalance) and the Matthews Correlation Coefficient (MCC), a balanced measure over all four "
       "confusion-matrix categories.")
bullet("Credit track: the Kolmogorov-Smirnov (KS) statistic, measuring the maximum separation "
       "between good and bad borrower score distributions (KS > 40% indicates a strong model), and "
       "the Gini coefficient, computed as Gini = 2*AUC - 1.")
body(
    "Every experiment logs its configuration and full metric set to a results table, enabling a "
    "consistent cross-model comparison in the Testing phase."
)

h2("3.5 Tools and Environment")
body(
    "The implementation uses Python with PyTorch and PyTorch Geometric for graph models, and "
    "scikit-learn / XGBoost for baselines. The repository follows a structured layout separating "
    "data, reusable source code (src/), exploratory notebooks, configuration files and reports. "
    "Reusable logic is kept in src/ while notebooks are used only for exploration, supporting "
    "reproducibility."
)

# =====================================================================
# 4. WORK DONE SO FAR
# =====================================================================
h1("4. Work Completed So Far")
body(
    "As of mid-June 2026, the Dissertation Outline phase is complete and the Design & Development "
    "phase is substantially advanced. The following items have been delivered."
)

prog = [
    ("Activity", "Status", "Notes"),
    ("Literature review", "Complete", "Surveyed GNN, temporal and autoencoder methods; identified the research gap."),
    ("Dataset selection", "Complete", "Tracks and datasets finalised; sources and licences documented."),
    ("Project scaffolding", "Complete", "Repository structure, environment and evaluation interface set up."),
    ("Data preprocessing & EDA", "In progress", "Cleaning, feature engineering and exploratory analysis for selected datasets."),
    ("Graph construction", "In progress", "Transaction/entity graphs built for Elliptic; heterogeneous graph for IEEE-CIS underway."),
    ("Baseline models", "Complete", "Logistic Regression, Random Forest and gradient boosting trained and scored."),
    ("GNN models", "In progress", "GCN and GraphSAGE implemented; GAT and tuning ongoing."),
    ("TCN / Autoencoder", "Started", "Architectures drafted; full training scheduled in the Testing phase."),
    ("Evaluation harness", "Complete", "Shared metrics module (PR-AUC, MCC, KS, Gini + standard) implemented."),
]
ptbl = doc.add_table(rows=0, cols=3)
ptbl.style = "Light Grid Accent 1"
for i, (a, b, c) in enumerate(prog):
    cells = ptbl.add_row().cells
    for cell, txt in zip(cells, (a, b, c)):
        para = cell.paragraphs[0]
        run = para.add_run(txt)
        run.font.size = Pt(9.5)
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, "1F3B73")
add_para("", space_after=6)

h2("4.1 Preliminary Findings")
body(
    "Early results are consistent with the literature. On the tabular fraud benchmark, gradient "
    "boosting forms a strong baseline, clearly outperforming Logistic Regression on PR-AUC under "
    "heavy imbalance. On the graph-native Elliptic data, initial GraphSAGE and GCN runs show that "
    "incorporating neighbourhood information improves recall on the illicit class relative to "
    "treating nodes independently, supporting the hypothesis that relational structure carries "
    "predictive signal. These numbers are preliminary and not yet tuned; the full, tuned "
    "comparison across all metrics is the objective of the upcoming Testing phase."
)

# =====================================================================
# 5. REMAINING PLAN
# =====================================================================
h1("5. Plan for Remaining Work")
plan = [
    ("Phase", "Timeline", "Planned work"),
    ("Design & Development (completion)", "By 20 June 2026", "Finalise heterogeneous graphs, complete GAT, TCN and autoencoder implementations."),
    ("Testing", "21 June - 10 July 2026", "Train, tune and evaluate all models; full cross-model comparison; result analysis."),
    ("Dissertation Review", "11 - 20 July 2026", "Submit draft to supervisor and examiner; incorporate feedback."),
    ("Submission", "21 - 28 July 2026", "Final corrections, documentation and dissertation submission."),
]
pl = doc.add_table(rows=0, cols=3)
pl.style = "Light Grid Accent 1"
for i, (a, b, c) in enumerate(plan):
    cells = pl.add_row().cells
    for cell, txt in zip(cells, (a, b, c)):
        para = cell.paragraphs[0]
        run = para.add_run(txt)
        run.font.size = Pt(9.5)
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, "1F3B73")
add_para("", space_after=6)

h2("5.1 Risks and Mitigation")
bullet("Scale of graph training (Elliptic / IEEE-CIS): mitigate with neighbour sampling "
       "(GraphSAGE-style mini-batching) and feature reduction.")
bullet("Severe class imbalance: mitigate with weighted losses, resampling and imbalance-aware "
       "metrics (PR-AUC, MCC) rather than accuracy.")
bullet("Fair comparison: enforce identical train/validation/test splits and a single evaluation "
       "harness across all models.")

# =====================================================================
# 6. CONCLUSION
# =====================================================================
h1("6. Conclusion")
body(
    "The dissertation is on schedule. The literature review and dataset selection are complete, "
    "the core infrastructure - repository, graph construction and a shared evaluation harness - is "
    "in place, baseline models are trained, and the first GNN models are running with encouraging "
    "preliminary results. The remaining effort is concentrated on completing the temporal and "
    "autoencoder models and on the rigorous, metric-driven comparison that will answer the central "
    "research question: whether graph- and sequence-aware deep learning meaningfully outperforms "
    "traditional machine learning for fraud detection and credit-default prediction."
)

# =====================================================================
# REFERENCES
# =====================================================================
h1("References")
refs = [
    "Cheng, D., Zou, Y., Xiang, S., and Jiang, C. “Review of Graph Neural Networks for "
    "Financial Fraud Detection.” arXiv preprint arXiv:2411.05815. "
    "Available at: https://arxiv.org/pdf/2411.05815",
    "Polu, O. R., Chamarthi, B., Chowdhury, T., Ushmani, A., Kasralikar, P., Syed, A. A., Mishra, "
    "A., Anumula, S. K., Rajendran, R. N., Mohanty, M. R., and Prova, N. N. I. “Graph Neural "
    "Networks for Fraud Detection: Modeling Financial Transaction Networks at Scale.”",
    "Weber, M., Domeniconi, G., Chen, J., et al. “Anti-Money Laundering in Bitcoin: "
    "Experimenting with Graph Convolutional Networks for Financial Forensics” (Elliptic "
    "dataset).",
    "Hamilton, W. L., Ying, R., and Leskovec, J. “Inductive Representation Learning on Large "
    "Graphs (GraphSAGE).” Advances in Neural Information Processing Systems, 2017.",
    "Kipf, T. N., and Welling, M. “Semi-Supervised Classification with Graph Convolutional "
    "Networks (GCN).” International Conference on Learning Representations, 2017.",
    "Velickovic, P., Cucurull, G., Casanova, A., et al. “Graph Attention Networks (GAT).” "
    "International Conference on Learning Representations, 2018.",
    "Bai, S., Kolter, J. Z., and Koltun, V. “An Empirical Evaluation of Generic Convolutional "
    "and Recurrent Networks for Sequence Modeling (TCN).” arXiv:1803.01271, 2018.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    p.add_run(f"[{i}] ").bold = True
    p.add_run(r).font.size = Pt(10.5)

# ---------- page numbers in footer ----------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
b1 = OxmlElement("w:fldChar"); b1.set(qn("w:fldCharType"), "begin")
it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
e1 = OxmlElement("w:fldChar"); e1.set(qn("w:fldCharType"), "end")
run._r.append(b1); run._r.append(it); run._r.append(e1)

out = "Mid_Semester_Report_Chandan_Nagavolu.docx"
doc.save(out)
print("Saved:", out)
